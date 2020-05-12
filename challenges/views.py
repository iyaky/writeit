import os, io
from django.shortcuts import render, redirect, get_object_or_404
import requests
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseRedirect, HttpResponse
from docx import Document
from .forms import PlainTextChallengeForm, FileChallengeForm, FilterForm
from django.contrib.auth.models import User
from challenges.models import Challenge, ChallengeTopicName
from register.models import Profile
from badges.models import Badge
from notification.models import Notification, NotificationType
from django.db.models import F
from django.core.files.storage import FileSystemStorage
from django.core.files.base import ContentFile
from django.core.files import File
from .filters import ChallengeFilter
import datetime

@login_required
def challenges(request):
    all_challenges = Challenge.objects.all()
    challenge_filter = ChallengeFilter(request.GET, queryset=all_challenges)
    return render(request, 'challenges/challenges.html', {'filter': challenge_filter})
    # if request.method == 'GET':
    #     form = FilterForm(request.GET)
    # else:
    #     form = FilterForm()
    # return render(request, 'challenges/challenges.html', {'form': form, 'challenges': all_challenges})

@login_required
def choose_challenge_type(request):
    return render(request, 'challenges/choose_challenge_type.html')


def auto_correct_text(text):
    r = requests.post("https://api.grammarbot.io/v2/check", data = {'text': text})
    j = r.json()
    new_text = ''
    cursor = 0
    l = 0;
    for match in j["matches"]:
        offset = match["offset"]
        length = match["length"]
        l += length
        if cursor > offset:
            continue
     # build new_text from cursor to current offset
        new_text += text[cursor:offset]
     # next add first replacement
        repls = match["replacements"]
        if repls and len(repls) > 0:
            new_text += repls[0]["value"]
     # update cursor
        cursor = offset + length

 # if cursor < text length, then add remaining text to new_text
    if cursor < len(text):
        new_text += text[cursor:]
    print("djcn")
    print(text)
    error_rate = l / len(text) * 100
    return {'new_text': new_text, 'error_rate': error_rate}


@login_required
def plain_text_challenge(request):
    if request.method == 'POST':
        form = PlainTextChallengeForm(request.POST)
        if form.is_valid():

            # save new challenge
            challenge = form.save()
            challenge.challenge_starter=request.user
            challenge.save()
            # manage badges and notifications
            challenges = Challenge.objects.filter(challenge_starter=request.user)
            count = 0
            for challenge in challenges:
                count += 1

            badges1 = Badge.objects.get(badge_name="Started 1 challenge")
            badges5 = Badge.objects.get(badge_name="Started 5 challenges")
            badges20 = Badge.objects.get(badge_name="Started 20 challenges")
            badges50 = Badge.objects.get(badge_name="Started 50 challenges")
            badges100 = Badge.objects.get(badge_name="Started 100 challenges")
            if count == 1 and not request.user in badges1.badge_users.all():
                badges1.badge_users.add(request.user)
                badges1.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 5 and not request.user in badges5.badge_users.all():
                badges5.badge_users.add(request.user)
                badges5.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 20 and not request.user in badges20.badge_users.all():
                badges20.badge_users.add(request.user)
                badges20.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 50 and not request.user in badges50.badge_users.all():
                badges50.badge_users.add(request.user)
                badges50.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 100 and not request.user in badges100.badge_users.all():
                badges100.badge_users.add(request.user)
                badges100.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)

            challenge.save()

            # first text auto check
            text = challenge.challenge_text
            auto_text_results = auto_correct_text(text)
            auto_text = auto_text_results['new_text']
            error_rate = auto_text_results['error_rate']
            challenge.auto_check_text = auto_text
            challenge.peer_review_text = auto_text

            # set status
            if challenge.number_of_checks == 0:
                challenge.status = 'Completed'
            else:
                challenge.status = "Pending"

            # set level
            if error_rate < 20:
                challenge.level = "Beginner"
            elif error_rate >= 20 and error_rate < 50:
                challenge.level = "Intermediate"
            elif error_rate >= 50 and error_rate < 80:
                challenge.level = "Advanced"
            else:
                challenge.level = "Expert"
            challenge.save()

            # second text auto check
            text = challenge.auto_check_text
            auto_text_results = auto_correct_text(text)
            auto_text = auto_text_results['new_text']
            challenge.auto_check_text = auto_text
            challenge.peer_review_text = auto_text
            challenge.save()

            # check if any brownie points available
            profile = Profile.objects.get(user = request.user)
            if profile.peer_review_points >= 1:
                challenge.completed_other_challenge = True
                challenge.save()
                profile.peer_review_points = F('peer_review_points') - 1
                profile.save()

                # manage notifications
                type = NotificationType.objects.get(name= "Create new challenge")
                new_notification(request, request.user.id, type.id)

            return redirect('peer_review_others', challenge_id=challenge.id)
    else:
        form = PlainTextChallengeForm()

    return render(request, 'challenges/plain_text_challenge.html', {'form': form})

@login_required
def file_challenge(request):
    if request.method == 'POST':
        form = FileChallengeForm(request.POST, request.FILES)
        if form.is_valid():

            # save new challenge
            challenge = form.save()
            challenge.challenge_starter=request.user

            # manage badges and notifications
            challenges = Challenge.objects.filter(challenge_starter=request.user)
            count = 0
            for i in challenges:
                count += 1

            badges1 = Badge.objects.get(badge_name="Started 1 challenge")
            badges5 = Badge.objects.get(badge_name="Started 5 challenges")
            badges20 = Badge.objects.get(badge_name="Started 20 challenges")
            badges50 = Badge.objects.get(badge_name="Started 50 challenges")
            badges100 = Badge.objects.get(badge_name="Started 100 challenges")
            if count == 1 and not request.user in badges1.badge_users.all():
                badges1.badge_users.add(request.user)
                badges1.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 5 and not request.user in badges5.badge_users.all():
                badges5.badge_users.add(request.user)
                badges5.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 20 and not request.user in badges20.badge_users.all():
                badges20.badge_users.add(request.user)
                badges20.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 50 and not request.user in badges50.badge_users.all():
                badges50.badge_users.add(request.user)
                badges50.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)
            elif count == 100 and not request.user in badges100.badge_users.all():
                badges100.badge_users.add(request.user)
                badges100.save()
                # manage notifications
                type = NotificationType.objects.get(name= "New badge")
                new_notification(request, request.user.id, type.id)

            challenge.save()

            file = Challenge.objects.get(pk=challenge.id)
            if file.uploaded_file:
                name, extension = os.path.splitext(file.uploaded_file.name)

                if extension == '.txt':
                    # open .txt file
                    file.uploaded_file.open('r')
                    text = file.uploaded_file.read()
                    file.uploaded_file.close()
                    text = str(text)

                    # first text auto check
                    challenge.challenge_text = text
                    auto_text_results = auto_correct_text(text)
                    auto_text = auto_text_results['new_text']
                    error_rate = auto_text_results['error_rate']
                    challenge.auto_check_text = auto_text
                    challenge.peer_review_text = auto_text
                    challenge.file_challenge = True

                    # set status
                    if challenge.number_of_checks == 0:
                        challenge.status = 'Completed'
                    else:
                        challenge.status = "Pending"

                    # set level
                    if error_rate < 20:
                        challenge.level = "Beginner"
                    elif error_rate >= 20 and error_rate < 50:
                        challenge.level = "Intermediate"
                    elif error_rate >= 50 and error_rate < 80:
                        challenge.level = "Advanced"
                    else:
                        challenge.level = "Expert"
                    challenge.save()

                    # second text auto check
                    text = challenge.auto_check_text
                    auto_text_results = auto_correct_text(text)
                    auto_text = auto_text_results['new_text']
                    challenge.auto_check_text = auto_text
                    challenge.peer_review_text = auto_text

                    # save .txt file with results
                    challenge.result_file.save('results.txt', ContentFile(auto_text))
                    challenge.save()

                elif extension == '.docx':
                    # open .docx file and read text to string
                    file_path = file.uploaded_file.path
                    doc = Document(file.uploaded_file)
                    text = '\n\n'.join(paragraph.text for paragraph in doc.paragraphs)
                    text = str(text)

                    # first text auto check
                    challenge.challenge_text = text
                    auto_text_results = auto_correct_text(text)
                    auto_text = auto_text_results['new_text']
                    error_rate = auto_text_results['error_rate']
                    challenge.auto_check_text = auto_text
                    challenge.peer_review_text = auto_text
                    challenge.file_challenge = True

                    # set status
                    if challenge.number_of_checks == 0:
                        challenge.status = 'Completed'
                    else:
                        challenge.status = "Pending"

                    # set level
                    if error_rate < 20:
                        challenge.level = "Beginner"
                    elif error_rate >= 20 and error_rate < 50:
                        challenge.level = "Intermediate"
                    elif error_rate >= 50 and error_rate < 80:
                        challenge.level = "Advanced"
                    else:
                        challenge.level = "Expert"
                    challenge.save()

                    # second text auto check
                    text = challenge.auto_check_text
                    auto_text_results = auto_correct_text(text)
                    auto_text = auto_text_results['new_text']
                    challenge.auto_check_text = auto_text
                    challenge.peer_review_text = auto_text

                    # save .docx file with results
                    result_doc = Document()
                    result_doc.add_paragraph(auto_text)

                    doc_io = io.BytesIO()
                    result_doc.save(doc_io)
                    doc_io.seek(0)

                    # Save the BytesIO to the field here
                    challenge.result_file.save("results.docx", File(doc_io))

                    response = HttpResponse(doc_io.read())
                    response["Content-Disposition"] = "attachment; filename=generated_doc.docx"
                    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                    challenge.save()

            else:
                print("File not found")

            # check if any brownie points available
            profile = Profile.objects.get(user = request.user)
            if profile.peer_review_points >= 1:
                challenge.completed_other_challenge = True
                challenge.save()
                profile.peer_review_points = F('peer_review_points') - 1
                profile.save()
                # manage notifications
                type = NotificationType.objects.get(name= "Create new challenge")
                new_notification(request, request.user.id, type.id)

            return redirect('peer_review_others', challenge_id=challenge.id)
    else:
        form = FileChallengeForm()

    return render(request, 'challenges/file_challenge.html', {'form': form})

@login_required
def peer_review_others(request, challenge_id):
    challenge = get_object_or_404(Challenge, pk=challenge_id)
    return render(request, 'challenges/peer_review_others.html', {'challenge': challenge})

@login_required
def challenge_detail(request, challenge_id):
    challenge_detail = get_object_or_404(Challenge, pk=challenge_id)
    now = datetime.date.today()
    dt = challenge_detail.deadline - now
    urgent = False
    if dt.days <= 3:
        urgent = True
    return render(request, 'challenges/challenge_detail.html', {'challenge': challenge_detail, "urgent": urgent})

@login_required
def challenge_accepted(request, challenge_id):
    if request.method == 'POST':
        challenge = get_object_or_404(Challenge, pk=challenge_id)
        challenge.peer_reviewer.add(request.user)
        challenge.agreed_number_of_checks = F('agreed_number_of_checks') + 1
        challenge.status = "Started"
        challenge.save()

        # manage notifications
        type = NotificationType.objects.get(name= "Accept new challenge")
        new_notification(request, request.user.id, type.id)

        return redirect('challenges')


    return render(request, 'challenges/challenges.html')


@login_required
def new_notification(request, user_id, notification_id):
    not_user = User.objects.get(pk=user_id)
    type = NotificationType.objects.get(pk=notification_id)
    not_type = type.id
    not_text = type.not_text
    new_notification = Notification.objects.create(to_user=not_user, type=type, notification_text=not_text)
    new_notification.save()
    return {'notification': new_notification}
